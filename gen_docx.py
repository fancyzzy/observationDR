#!/usr/bin/env python3

'''
写docx
'''

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docx.shared import Mm
from docx.shared import Inches, Pt

#样式，字体,table宽度高度
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import os
from datetime import datetime
from collections import namedtuple
from copy import deepcopy
from numpy import array
from numpy import isnan
from numpy import nan
import read_xlsx

from my_log import printl

#防止matplotlib在多线程调用下
#造成Tkinter主线程crash
from matplotlib import use
use('Agg')
import draw_plot

ProInfo = namedtuple("ProInfo", ['name', 'area', 'code', 'contract', 'builder',\
		'supervisor', 'third_observer', 'builder_observer', 'xlsx_path', 'date'])



#日报信息头页，总体监测分析表， 现场巡查表， 沉降监测表(地表，建筑物，管线),
#测斜监测表，爆破振动监测表，平面布点图
PAGE_CATEGORY = ['header', 'overview', 'security', 'settlement_ground',\
	'settlement_buidling', 'settlement_pipeline', 'inclinometer', 'blasting',\
	'floor_layout']


def date_to_str(date_str):
	ds = date_str.strftime("%Y/%m/%d")
	return ds.split('/')[0] + '年' + ds.split('/')[1].lstrip('0')\
	 + '月' + ds.split('/')[2].lstrip('0') + '日'

def d2s(date_str):
	ds = date_str.strftime("%Y/%m/%d")
	return ds.split('/')[1] + '月' + ds.split('/')[2] + '日'

def d_s(date_str):
	ds = date_str.strftime("%Y/%m/%d")
	return ds

def get_file_list(dir,file_list):
	'''
	获取目录dir下的所有文件名(文件路径)
	略过隐藏的特殊文件
	支持子目录
	'''
	try:
		new_dir = dir
		if os.path.isfile(dir):
			file_list.append(dir)
		elif os.path.isdir(dir):
			for s in os.listdir(dir):
				#略过特殊字符开头的文件或者文件夹
				if not s[0].isdigit() and not s[0].isalpha():
					#logger.warning("Hidden file:%s"%(s))
					#logger.warning("Hidden file:{}".format(s))
					if s != '.':
						continue
				new_dir = os.path.join(dir,s)
				get_file_list(new_dir,file_list)
		else:
			pass
	except Exception as e:
		#logger.warning(e)
		print("warning,e:",e)
	return file_list


def get_max(l):
	'''
	兼容有numpy.nan或者None的array和list
	'''
	max_v = 0
	for item in l:

		if item != None and item != nan:
			if item > max_v:
				max_v = item
	return max_v


def delete_item(item):
	'''
	删除docx中的元素，table，paragraph
	'''
	ie = item._element
	ie.getparent().remove(ie)
	ie._p = ie._element = None



class MyDocx(object):
	def __init__(self, docx_path, proj_info, my_xlsx, alarm_feature=True):
		self.proj = ProInfo(*proj_info)
		self.docx = None
		self.path = docx_path
		self.date = proj_info[-1]
		self.xlsx_path = os.path.dirname(proj_info[-2])
		self.str_date = date_to_str(self.date)
		#xlsx实例
		self.my_xlsx = my_xlsx
		self.my_plot = draw_plot.MyPlot()
		#报警实例
		self.alarm_feature = alarm_feature
		self.my_alarm = MyAlarm()
		self.alarm_r = None

		#签名文件列表
		self.sig_list = []
		sig_path = os.path.join(self.xlsx_path,'签名')
		item_list = os.listdir(sig_path)
		for item in item_list:
			if '.png' in item or '.PNG' in item:
				self.sig_list.append(os.path.join(sig_path, item))
		print("DEBUG self.sig_list=",self.sig_list)

	#########__init__()#####################################

	def get_table_num(self):
		return len(self.docx.tables)

	###############get_table_num###########################


	def gen_docx(self):
		'''
		生成docx文件
		'''
		printl("\ngen_docx:{}日报:".format(self.str_date),False)

		#if not self.path or not os.path.exists(self.path):
		if not self.path:
			printl("error, no available docx path")
			return
		
		#读取'default_template.docx'
		self.docx = Document()
		self.set_document_style()

		#页面布局为A4 宽210mm*高297mm
		section = self.docx.sections[0]
		section.page_width = Mm(210)
		section.page_height = Mm(297)

		#首页
		printl("\n###1. 报表首页###")
		if not self.make_header_pages():
			printl("DEBUG make_head_page error")
		else:
			printl("1@ 生成首页")
			pass


		#数据分析表****
		#12 percentage
		printl("\n###2. 监测数据分析表###")
		if not self.make_overview_pages():
			printl("DEBUG make_overview_pages error")
		else:
			self.docx.save(self.path)



		#页面布局为横向
		new_section = self.docx.add_section(WD_SECTION.NEW_PAGE)
		new_section.orientation = WD_ORIENT.LANDSCAPE
		new_section.page_width = Mm(297)
		new_section.page_height = Mm(210)
		new_section.top_margin = Cm(2.8)
		new_section.bottom_margin = Cm(2.6)
		new_section.left_margin = Cm(2.5)
		new_section.right_margin = Cm(2.7)
		new_section.header_distance = Cm(1)
		new_section.footer_distance = Cm(1)

		#现场安全巡视页
		#2 percentage
		printl("\n###3. 现场安全巡视表###")
		if not self.make_security_pages():
			printl("DEBUG make_security_pages error")
		else:
			pass

		#页面布局为纵向横向
		new_section = self.docx.add_section()
		new_section.orientation = WD_ORIENT.PORTRAIT
		new_section.page_width = Mm(210)
		new_section.page_height = Mm(297)
		new_section.top_margin = Cm(2.7)
		new_section.bottom_margin = Cm(2.5)
		new_section.left_margin = Cm(2.8)
		new_section.right_margin = Cm(2.6)
		new_section.header_distance = Cm(1)
		new_section.footer_distance = Cm(1)

		#沉降监测表页
		#45 percent in pages
		printl("\n###4. 沉降监测报表###")
		self.allow_page_break = False
		if not self.make_settlement_pages():
			printl("DEBUG make_settlement_pages error")
		else:
			self.docx.save(self.path)

		#测斜监测表页
		#35 percent in pages
		printl("\n###5. 测斜监测报表###")
		if not self.make_inclinometer_pages():
			printl("DEBUG make_inclinometer_pages error")
		else:
			self.docx.save(self.path)

		#new section landscape
		#页面布局为横向
		new_section = self.docx.add_section(WD_SECTION.NEW_PAGE)
		new_section.orientation = WD_ORIENT.LANDSCAPE
		new_section.page_width = Mm(297)
		new_section.page_height = Mm(210)
		new_section.top_margin = Cm(2.17)
		new_section.bottom_margin = Cm(2.17)
		new_section.left_margin = Cm(2.54)
		new_section.right_margin = Cm(2.54)
		new_section.header_distance = Cm(1.5)
		new_section.footer_distance = Cm(1.75)

		#爆破振动监测报表
		printl("\n###6. 爆破振动监测报表###")
		if not self.make_blasting_pages():
			printl("DEBUG make_blasting_pages error")
		else:
			printl("1@ 生成爆破监测表")
			pass

		#页面布局
		new_section = self.docx.add_section(WD_SECTION.NEW_PAGE)
		new_section.orientation = WD_ORIENT.LANDSCAPE
		new_section.page_width = Mm(297)
		new_section.page_height = Mm(210)
		new_section.top_margin = Cm(3.17)
		new_section.bottom_margin = Cm(3.17)
		new_section.left_margin = Cm(2.54)
		new_section.right_margin = Cm(2.54)
		new_section.header_distance = Cm(1.5)
		new_section.footer_distance = Cm(1.75)


		#平面布点图表
		printl("\n###7. 平面布点图###")
		if not self.make_layout_pages():
			printl("DEBUG make_layout_pages error")
		else:
			printl("2@ 生成平面布点图")
			pass

		#alarm
		if self.alarm_r and self.my_alarm.alarm_on:
			self.alarm_r.text = "报        警： 是 √         否"

		#保存
		self.docx.save(self.path)
		printl("日报生成结束!")
		printl("saved in:'{}'".format(self.path))
		return True
	#######gen_docx()####################################

	def set_document_style(self):
		'''
		设定全局字体样式
		'''
		d = self.docx
		d.styles['Normal'].font.name = 'Times New Roman'
		r = d.styles['Normal']._element
		r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
	############set_document_style()####################

	def write_header(self):
		'''
		项目信息
		'''
		d = self.docx
		styles = d.styles
		p = d.add_paragraph()
		r = p.add_run(self.proj.name)
		r.font.size = Pt(18)
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		p.paragraph_format.space_before = Pt(32)
		p.paragraph_format.space_after = 0
		p.paragraph_format.line_spacing_rule = 1
		p.paragraph_line_spacing = None


		#add a new style 
		new_style = d.styles.add_style('my_sub_header', WD_STYLE_TYPE.PARAGRAPH)
		p_format = new_style.paragraph_format
		p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
		p_format.space_after = 0
		p_format.space_before = 0
		p_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
		p_format.line_spacing = Pt(21)

		font = new_style.font
		font.name = 'Times New Roman'
		r = new_style._element
		r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
		font.size = Pt(12)

		p = d.add_paragraph("施工单位:  ")
		p.style = styles['my_sub_header']
		p.add_run("%s" % self.proj.builder).underline = True
		p.add_run("    合同号：")
		p.add_run("%s" % self.proj.contract).underline = True

		p = d.add_paragraph("监理单位：")
		p.style = styles['my_sub_header']
		p.add_run("%s" % self.proj.supervisor).underline = True
		p.add_run("               编号：")
		p.add_run("%s" % self.proj.code).underline = True

		p = d.add_paragraph("第三方检测单位：")
		p.style = styles['my_sub_header']
		p.add_run("%s" % self.proj.third_observer).underline = True
	################write_header()########################


	def make_header_pages(self):
		'''
		首页
		'''
		result = False
		d = self.docx

		styles = d.styles

		p = d.add_paragraph()
		r = p.add_run(self.proj.name)
		r.bold = True
		r.font.size = Pt(18)
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		p.paragraph_format.space_before = Pt(32)
		p.paragraph_format.space_after = 0
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.paragraph_line_spacing = None

		p = d.add_paragraph()
		r = p.add_run(self.proj.area)
		r.underline = True
		r.font.size = Pt(16)
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.paragraph_format.space_after = 0
		p.paragraph_format.line_spacing = Pt(28)

		p = d.add_paragraph()
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		p.line_spacing_rule = WD_LINE_SPACING.EXACTLY
		p.paragraph_format.line_spacing = Pt(28)

		p = d.add_paragraph()
		r = p.add_run("第三方检测日报")
		r.bold = True
		r.font.size = Pt(22)
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		p.paragraph_format.space_before = 0
		p.paragraph_format.space_after = 0
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


		p = d.add_paragraph()
		r = p.add_run("(第%s次)"%self.proj.code.split('-')[-1])
		r.font.size = Pt(16)
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
		p.paragraph_format.line_spacing = Pt(28)


		#增加一个style
		new_style = d.styles.add_style('my_header', WD_STYLE_TYPE.PARAGRAPH)
		p_format = new_style.paragraph_format
		p_format.first_line_indent = Cm(4.2)
		p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
		p_format.line_spacing = Pt(28)
		p_format.space_after = 0
		p_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
		font = new_style.font
		font.name = 'Times New Roman'
		r = new_style._element
		r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
		font.size = Pt(15)

		p = d.add_paragraph(style = styles['my_header'])

		p = d.add_paragraph()
		p.style = styles['my_header']
		r = p.add_run("编        号：")
		r = p.add_run("%s"%self.proj.code)
		r.underline = True

		p = d.add_paragraph("检测日期：")
		p.style = styles['my_header']
		p.add_run("%s"%self.str_date).underline = True

		p = d.add_paragraph()
		p.style = styles['my_header']

		p = d.add_paragraph()
		p.style = styles['my_header']
		self.alarm_r = p.add_run()
		self.alarm_r.text = "报        警： 是         否 √"

		p = d.add_paragraph()
		p.style = styles['my_header']
		p.add_run("报警内容:  ")

		for i in range(3):
			d.add_paragraph(style=styles['my_header'])

		p = d.add_paragraph("项目负责人：")
		p.style = styles['my_header']
		p.add_run(" "*20+".").underline = True

		p = d.add_paragraph("签发日期：")
		p.style = styles['my_header']
		p.add_run(" "*22+".").underline = True

		p = d.add_paragraph("单位名称：")
		p.style = styles['my_header']
		p.add_run("      (盖章)     .").underline = True

		d.add_paragraph(style = styles['my_header'])

		p = d.add_paragraph()
		r = p.add_run(self.str_date)
		r.font.size = Pt(15)
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		###new page###########
		#审核意见单
		d.add_page_break()
		self.write_header()

		p = d.add_paragraph()
		r = p.add_run("第三方检测审核单")
		r.bold = True
		r.font.size = Pt(18)
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
		p.paragraph_format.line_spacing = Pt(21)
		p.paragraph_format.space_before = Pt(21)

		t = d.add_table(rows=1, cols=1, style = 'Table Grid')
		t.alignment = WD_TABLE_ALIGNMENT.CENTER
		tr = t.rows[0]
		tr.herght = Inches(6)
		tr.width = Inches(6)

		p = t.cell(0,0).paragraphs[0]
		r = p.add_run("审核意见：  ")
		r.font.size = Pt(14)

		for i in range(8):
			p = t.cell(0,0).add_paragraph()
			r = p.add_run()
			r.font.size = Pt(14)
			p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
			p.paragraph_format.line_spacing = Pt(21)

		p = t.cell(0,0).add_paragraph()
		s = " "*60 +"监理工程师：" + " "*20 + "日期：" 
		r = p.add_run(s)
		r.font.size = Pt(12)

		for i in range(3):
			p = t.cell(0,0).add_paragraph()
			p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
			p.paragraph_format.line_spacing = Pt(21)

		result = True
		return result
	##################make_header_pages()################	

	def get_values_by_field(self, sheet, row_list, field_name):
		'''
		根据要获取的列域名，得到row_list的行的数值
		'''
		px = self.my_xlsx
		value_list = []
		field_col,_ = px.get_item_point(sheet, field_name)
		if field_col:
			#print("DEBUG sheet:{}, row_list={}, field_col={}".format(sheet, row_list, field_col))
			value_list = px.get_rows_col_values(sheet,row_list,field_col)

		return value_list

	#######get_values_by_field()###########################


	def get_value_by_field(self, sheet, area_name, field_name):
		'''
		根据要获取的列域名，获取area_name的第一行的值
		'''
		px = self.my_xlsx
		field_value = ''

		field_col,field_row = px.get_item_point(sheet, field_name)
		if field_col and field_row:
			row_start, row_end = px.all_areas_row_range[sheet][area_name]
			field_value = px.get_value(sheet, row_start, field_col)

		if field_value == None:
			field_value = ' '
		#print("DEBUG '{}'在{}:{}的值为:{}".format(field_name,sheet,area_name,field_value))

		return field_value
	################get_value_by_field()##################


	def get_col_values(self, sheet, area_name, col, d_obser_range, obser_list):
		'''
		获取指定列的值
		根据不同的sheet，有不同的算法
		比如倾斜会求每两排的差值
		支撑轴力会求平方差再乘以系数
		但是都只有一个值是动态，根据col的不同

		d_obser_range :是点号的行范围字典,用于一个点号对应多行值，需要求平均的sheet
		'''
		def is_number(s):
			try:
				float(s)
				return True
			except:
				pass
			return False

		px = self.my_xlsx
		sh = px.wb[sheet]
		output_values = px.get_range_values(sheet, area_name, col)
		ln = len(output_values)
		try:
			output_values = array(output_values, dtype=float)
		except:
			print("DEBUG,有非数值value: {}在sheet: {}, area_name: {}, col: {}".\
				format(output_values,sheet,area_name,col))
			for i in range(ln):
				if not is_number(output_values[i]):
					output_values[i] = None
			output_values = array(output_values, dtype=float)

		#获取area的测量点行范围row_range
		start_row, end_row = px.all_areas_row_range[sheet][area_name]

		#建筑物倾斜的值每两个做差，然后把差值赋值回第一个，第二个设为None
		if '建筑物倾斜' in sheet:
			tmp_values = []
			for i in range(ln-1):
				if i%2 != 0:
					continue
				curr_v = output_values[i]
				next_v = output_values[i+1]
				tmp_values.append(next_v - curr_v)
				tmp_values.append(nan)
			output_values = array(tmp_values, dtype=float)

		elif '支撑轴力' in sheet:
			before_values = []
			col,_ = px.get_item_point(sheet, '埋设前', from_last_search=False)
			if col == None:
				print("Error,缺少埋设前列!")
				before_values = [None for x in range(ln)]
			else:
				for i in range(start_row, end_row+1):
					before_values.append(sh.cell(i, col).value)
			before_values = array(before_values, dtype=float)

			factor_values = []
			col,_ = px.get_item_point(sheet, '率定系数', from_last_search=False)
			if col == None:
				print("Error, 缺少率定系数列!")
				factor_values = [None for x in range(ln)]
			else:
				for i in range(start_row, end_row+1):
					factor_values.append(sh.cell(i, col).value)
			factor_values = array(factor_values, dtype=float)

			tmp_values = []
			#当天值和埋设前的平法差再乘以率定系数
			tmp_values = (output_values**2 - before_values**2) * factor_values

			output_values = tmp_values

		elif '混撑' in sheet:
			#算法, 单位mm所以要/10**9
			#(初值**2 - 埋设前or当天值**2)*率定系数*混凝土支撑截面积*弹性模量/10**9
			#求初值要加'-'
			#如果是求初始轴力，那么output_values应该是当天值换成埋设前
			initial_values = [] #初值
			factor_values = [] #率定系数
			area_values = [] #混凝土支撑截面积
			elastic_values = [] #弹性模量
			col,_ = px.get_item_point(sheet, '初值', from_last_search=False)
			if col == None:
				printt("Error, 缺少初值列")
				initial_values = [None for x in range(ln)]
			else:
				for i in range(start_row, end_row+1):
					initial_values.append(sh.cell(i, col).value)
			initial_values = array(initial_values, dtype=float)

			col,_ = px.get_item_point(sheet, '率定系数', from_last_search=False)
			if col == None:
				printt("Error, 缺少率定系数列")
				factor_values = [None for x in range(ln)]
			else:
				for i in range(start_row, end_row+1):
					factor_values.append(sh.cell(i, col).value)
			factor_values = array(factor_values, dtype=float)

			col,_ = px.get_item_point(sheet, '混凝土支撑截面积', from_last_search=False)
			if col == None:
				printt("Error, 缺少混凝土支撑截面积列")
				area_values = [None for x in range(ln)]
			else:
				for i in range(start_row, end_row+1):
					area_values.append(sh.cell(i, col).value)
			area_values = array(area_values, dtype=float)

			col,_ = px.get_item_point(sheet, '弹性模量', from_last_search=False)
			if col == None:
				printt("Error, 缺少弹性模量列")
				elastic_values = [None for x in range(ln)]
			else:
				for i in range(start_row, end_row+1):
					elastic_values.append(sh.cell(i, col).value)
			elastic_values = array(elastic_values, dtype=float)

			tmp_values = []
			tmp_values = (initial_values**2 - output_values**2)\
			 *factor_values *area_values *elastic_values /10**9

			output_values = tmp_values

			#求平均
			tmp_values = []
			#第一个观测点的行坐标为对照，用来获取对应的index
			base,_ = d_obser_range[obser_list[0]]
			start = 0
			end = 0
			for obser in obser_list:
				start, end = d_obser_range[obser]
				add_count = 0
				v_sum = 0
				v_average = nan
				for i in range(start-base, end-base+1):
					if not isnan(output_values[i]):
						v_sum += output_values[i]
						add_count += 1
				if add_count != 0:
					v_average = v_sum/(add_count)
				#只把第一个值赋值为平均值，其他为nan
				for i in range(start-base, end-base+1):
					if i == start-base:
						tmp_values.append(v_average)
					else:
						tmp_values.append(nan)

			output_values = array(tmp_values, dtype=float)
			#print("DEBUG 混撑平均值:",output_values)

		elif '锚索轴力' in sheet:
			#算法:
			#本次轴力： (埋设前**2 - date_value**2)*率定系数
			#初始轴力:  (埋设前**2 - 初始值**2)*率定系数
			#求平均，只保留点号第一行为平均值，其他都是None/nan
			before_values = [] #初值
			factor_values = [] #率定系数
			col,_ = px.get_item_point(sheet, '埋设前', from_last_search=False)
			if col == None:
				printt("Error, 缺少埋设前列")
				before_values = [None for x in range(ln)]
			else:
				for i in range(start_row, end_row+1):
					before_values.append(sh.cell(i, col).value)
			before_values = array(before_values, dtype=float)

			col,_ = px.get_item_point(sheet, '率定系数', from_last_search=False)
			if col == None:
				printt("Error, 缺少率定系数列")
				factor_values = [None for x in range(ln)]
			else:
				for i in range(start_row, end_row+1):
					factor_values.append(sh.cell(i, col).value)
			factor_values = array(factor_values, dtype=float)

			tmp_values = []
			tmp_values = (before_values**2 - output_values**2)* factor_values

			output_values = tmp_values

			#求平均
			tmp_values = []
			#第一个观测点的行坐标为对照，用来获取对应的index
			base,_ = d_obser_range[obser_list[0]]
			start = 0
			end = 0
			for obser in obser_list:
				start, end = d_obser_range[obser]
				add_count = 0
				v_sum = 0
				for i in range(start-base, end-base+1):
					if not isnan(output_values[i]):
						v_sum += output_values[i]
						add_count += 1
				v_average = v_sum/(add_count)
				#只把第一个值赋值为平均值，其他为nan
				for i in range(start-base, end-base+1):
					if i == start-base:
						tmp_values.append(v_average)
					else:
						tmp_values.append(nan)

			output_values = array(tmp_values, dtype=float)
			#print("DEBUG 锚索轴力:",output_values)


		return output_values
	###########get_col_values()##########################


	def get_diff_values(self, sheet, l_values, r_values):
		'''
		diff_values = (l-values - r_values)*1000
		根据不同的sheet，有不同的求差值的算法
		'''
		output_values = []

		if '建筑物倾斜' in sheet:
			output_values = (array(l_values, dtype=float) - \
			array(r_values, dtype=float))/15*1000
			#print("DEBUG diff_values: {} = ({} - {})/15*1000".\
			#	format(output_values,l_values,r_values))

		elif '支撑轴力' in sheet:
			output_values = (array(l_values, dtype=float) - \
			array(r_values, dtype=float))
			#print("DEBUG 支撑轴力，变化值:",output_values)

		elif '混撑' in sheet:
			output_values = (array(l_values, dtype=float) - \
			array(r_values, dtype=float))

		elif '锚索轴力' in sheet:
			output_values = (array(l_values, dtype=float) - \
			array(r_values, dtype=float))

		elif '净空收敛' in sheet:
			#unit is mm, therefore no 1000 multiplying
			output_values = (array(l_values, dtype=float) - \
			array(r_values, dtype=float))

		else:
			output_values = (array(l_values, dtype=float) - \
			array(r_values, dtype=float))*1000

		return output_values
	##########get_diff_values()###########################


	def get_acc_values(self, sheet, l_values, r_values, o_acc_values):
		'''
		acc_values = (l_values - r_values)*1000 + o_acc_values
		根据不同的sheet，有不同的求差值的算法
		'''
		output_values = []

		if '建筑物倾斜' in sheet:
			output_values = (array(l_values, dtype=float) - \
			array(r_values, dtype=float))/15*1000 + array(o_acc_values,\
			dtype=float)
			#print("DEBUG 建筑物倾斜 累计变化量 acc_values: {} = ({} - {})/15*1000 + {}".\
			#	format(output_values,l_values,r_values, o_acc_values))

		elif '支撑轴力' in sheet:
			output_values = (array(l_values, dtype=float) - \
			array(r_values, dtype=float)) + array(o_acc_values,\
			dtype=float)
			#print("DEBUG 支撑轴力 本次轴力:",l_values)
			#print("DEBUG 支撑轴力 初始轴力:",r_values)
			#print("DEBUG 支撑轴力 旧累计:",o_acc_values)
			#print("DEBUG 支撑轴力 累计变化量:",output_values)

		elif '混撑' in sheet:
			output_values = (array(l_values, dtype=float) - \
			array(r_values, dtype=float)) + array(o_acc_values,\
			dtype=float)

		elif '锚索轴力' in sheet:
			output_values = (array(l_values, dtype=float) - \
			array(r_values, dtype=float)) + array(o_acc_values,\
			dtype=float)

		elif '净空收敛' in sheet:
			output_values = (array(l_values, dtype=float) - \
			array(r_values, dtype=float)) + array(o_acc_values,\
			dtype=float)

		else:
			output_values = (array(l_values, dtype=float) - \
			array(r_values, dtype=float))*1000 + array(o_acc_values,\
			dtype=float)

		return output_values
	##########get_acc_values()###########################


	def one_overview_table(self, area_name, v_percent, per_s):
		'''
		一个区间的各种观测监信息汇总表
		'''
		d = self.docx
		px = self.my_xlsx
		#t = d.add_table(rows=1, cols=8, style='Table Grid')
		t = d.add_table(rows=1, cols=8, style='my_table_style')

		t.cell(0,0).text = '监测项目'
		t.cell(0,1).text = '本次\n变化\n最大点'
		t.cell(0,2).text = '日变化\n速率\n(mm/d)'
		t.cell(0,3).text = '日变量\n报警值\n(mm/d)'
		t.cell(0,4).text = '累计\n变化量\n最大点'
		t.cell(0,5).text = '累计\n变化量\n/mm'
		t.cell(0,6).text = '累计\n变量\n报警值/mm'
		t.cell(0,7).text = '累计\n变量\n控制值/mm'

		#找到这个area的所有相关的观测项目页
		related_sheets = []
		for sheet in px.sheets:
			if area_name in px.all_areas_row_range[sheet].keys():
					related_sheets.append(sheet)

		total_sheets_num = len(related_sheets)
		#记录写入有效行数
		count_num = 0
		print("{}个观测项目数据: {}".format(total_sheets_num, related_sheets))

		#遍历这个区间站所存在的观测页表格	
		sub_v_percent = v_percent/total_sheets_num
		for sheet in related_sheets:
			#if '锚索轴力' in sheet:
			#	print("DEBUG 锚索轴力 数据分析表")
			if '孔深测斜' in sheet:
				printl("%f@"%(sub_v_percent))
				continue

			printl("{}数据分析表:'{}'({}/{}:{})".format(per_s,\
				area_name, count_num, total_sheets_num, sheet))
			#获取数据
			today_range_values = [] #当天数据列表
			last_range_values = []  #前一天数据列表
			diff_original_values = [] #变化量数据列表
			diff_abs_values = []      #变化量绝对值数据列表

			col_alpha = 'B'
			obser_col,obser_row = px.get_item_point(sheet,'点号',from_last_search=False)
			if obser_col == None:
				obser_col,obser_row = px.get_item_point(sheet,'测点',from_last_search=False)
			if obser_col == 3:
				col_alpha = 'C'
			if obser_col == 4:
				col_alpha = 'D'
			if obser_col == None:
				printl("Error, {}无观测点列!".format(sheet))
				printl("%f@"%(sub_v_percent))
				continue

			#点号的行范围字典，适用于一个点号对应多行数值的sheet
			d_obser_range = {}
			obser_list = []
			if '混撑' in sheet or '锚索轴力' in sheet:
				#以点号/测点为锚点，找每个点号/测点的行范围
				d_obser_range = px.get_one_sheet_areas_range(sheet, obser_col, obser_row+1)
				start, end = px.all_areas_row_range[sheet][area_name]
				for i in range(start, end+1):
					#获取第二列点号的值
					obser_name = px.get_value(sheet, i, obser_col)
					if obser_name != None:
						obser_list.append(obser_name)

			#获取当天的数据
			#获取当天值这一列的坐标
			today_col,today_row = px.get_item_point(sheet, self.date)
			if today_col == None:
				printl("Warning, {}的区间{}没有当天值列!".format(sheet,area_name))
				printl("%f@"%(sub_v_percent))
				continue

			#获取当天数据列
			today_range_values = self.get_col_values(sheet, area_name, today_col, d_obser_range, obser_list)
			#print("DEBUG sheet:{}, today_range_values:{}".format(sheet,today_range_values))

			#如果所有值都为空就略过这一行的填写
			if isnan(today_range_values).sum() == len(today_range_values):
				printl("{}数据分析表:'{}'({}/{}:{}, 没有当天数据!)".format(per_s,\
				area_name, count_num, total_sheets_num, sheet))
				continue

			#寻找前一天数据
			last_date = px.get_value(sheet, today_row, today_col-1)
			#print("DEBUG 前一天列坐标:'{}',值:'{}',值类型'{}'".format(today_col-1,last_date,str(type(last_date))))
			if 'datetime' in str(type(last_date)):
				last_range_values = self.get_col_values(sheet, area_name, today_col-1,\
				 d_obser_range, obser_list) 
				if isnan(last_range_values).sum() == len(last_range_values):
					printl("没有前一天数据!")
					continue
			else:
			#前面一列不是日期值，表示昨天值不存在
				last_range_values = None
				#不存在就略过这一行的数据
				printl("Warning, 没有前一天数据值列!")
				continue
			#print("DEBUG 上一次值lastday_range_values:",last_range_values)

			#找到其中绝对值最大为变化最大的
			diff_original_values = self.get_diff_values(sheet, today_range_values,\
				last_range_values)
			#print("DEBUG 本次变化diff:",diff_original_values)

			#如果都是nan就略过这一行的数据填写 
			if isnan(diff_original_values).sum() == len(diff_original_values):
				printl("变化值全部为空!")
				continue

			#求出绝对值最大的值
			#负值求abs会造成小数点后面数字很多，用round(x,2)只保留2位有效数
			diff_abs_values = [round(abs(value),2) for value in diff_original_values]
			#print("DEBUGf diff_abs_values:",diff_abs_values)
			max_change = get_max(diff_abs_values)
			#print("DEBUG max_change='{}'".format(max_change))

			#列出所有max 点
			max_obser_list = []
			max_change_values = []
			row_list = []
			row_start, row_end = px.all_areas_row_range[sheet][area_name]
			if max_change != 0 and not isnan(max_change):
				for i, v in enumerate(diff_abs_values):
					if max_change == v:
						#找到区间的行范围, 加上最大值的相对index就是最大值的row_index
						row_index = i + row_start
						row_list.append(row_index)
						s_index = col_alpha + '%d'%row_index
						obser_id = px.wb[sheet][s_index].value
						max_obser_list.append(obser_id)
						max_value = str(round(diff_original_values[i],2))
						max_change_values.append(max_value)
					else:
						continue
				print("本次变化最大点:{},值:{}".format(\
					max_obser_list, max_change_values))
			else:
				printl("warning, 没有最大值!")
				#略过这一行的填写
				continue

			count_num += 1
			#新加一行，写入测量项目sheet，写入这个测量点id
			row = t.add_row()
			#监测项目
			row.cells[0].text = sheet


			#日变量报警值
			field_name = '日变量报警值'
			field_values = self.get_values_by_field(sheet, row_list, field_name)
			#print("DEBUG 日变量报警值:",field_values)
			s = ''
			if len(field_values) == 0: 
				s = ' '
			else:
				for value in field_values:
					if value == None:
						s += ' ' + '\n'
					else:
						s += str(value) + '\n'
			row.cells[3].text = s.strip('\n')

			'''
			ln = len(field_values)
			cell = row.cells[3]
			p_cell = row.cells[3].paragraphs[0]
			runs = []
			r = None
			if ln == 0: 
				cell.text = ' '
			else:
				for i in range(ln):
					if i != ln-1:	
						if field_values[i] == None:
							r = p_cell.add_run(' \n')
						else:
							r = p_cell.add_run(field_values[i] + '\n')
					else:
						if field_values[i] == None:
							r = p_cell.add_run(' ')
						else:
							r = p_cell.add_run(field_values[i])
						r.font.bold = True
					runs.append(r)

			print("DEBUGGGGG 这个单元格一共有多少个r:",len(runs))
			'''

			#本次变化最大点
			if not self.alarm_feature:
				s = ''
				for obser in max_obser_list:
					s += obser + '\n' 
				row.cells[1].text = s.strip('\n')
			else:
				self.set_cell_text_by_field_values(row.cells[1], max_obser_list, max_change_values, field_values)


			#日变化速率
			s = ''
			for max_v in max_change_values:
				s += max_v + '\n'
			row.cells[2].text = s.strip('\n')


			#求本次累计值 = 当前值-初值+旧累计值
			acc_values = [] #累计变化量列表
			acc_abs_values = [] #累计变化量绝对值列表
			#获取'初值'这一列，在第3列
			#如果是混撑，换成‘埋设前’值
			initial_name = '初值'
			if '混撑' in sheet:
				initial_name = '埋设前'

			init_col,init_row = px.get_item_point(sheet, initial_name)
			if init_col == None:
				printl("Error, sheet:{}没有{}列!".format(sheet,initial_name))
				continue
			initial_range_values = self.get_col_values(sheet, area_name, init_col,\
				d_obser_range, obser_list)

			if '混撑' in sheet:
				initial_range_values = -initial_range_values
				#printl("DEBUG 混撑，初始轴力平均值:{}".format(initial_range_values))

			#printl("DEBUG '初始值列':{}".format(initial_range_values))
			#获取'旧累计'这一列，在第4列
			old_acc_col,_ = px.get_item_point(sheet, '旧累计')
			if old_acc_col != None:
				old_acc_range_values = self.get_col_values(sheet, area_name, old_acc_col,\
					d_obser_range, obser_list)
				#处理旧累计，如果为None就设为0
				ln = len(old_acc_range_values)
				for i in range(ln):
						if isnan(old_acc_range_values[i]):
							old_acc_range_values[i] = 0
			else:
				#支撑轴力/混撑/锚索轴力不存在旧累计列
				old_acc_range_values = [0 for x in range(len(initial_range_values))]

			acc_values = self.get_acc_values(sheet,today_range_values,\
				initial_range_values, old_acc_range_values)

			if '混撑' in sheet:
				printl("DEBUG 混撑, 本次累计:{}".format(acc_values))

			#printl("DEBUG '本次累计值列':{}".format(acc_values))

			acc_abs_values = [round(abs(v),2) for v in acc_values]
			max_acc = get_max(acc_abs_values)

			#列出所有max 点
			max_obser_list = []
			max_acc_values = []
			row_list =[]
			#printl("DEBUG '最大累计变化值'是:{}".format(max_change))
			row_start, row_end = px.all_areas_row_range[sheet][area_name]
			if max_acc != 0 and not isnan(max_acc):
					for i, v in enumerate(acc_abs_values):
						if max_acc == v:
							#找到区间的行范围, 加上最大值的相对index就是最大值的row_index
							row_index = i + row_start
							row_list.append(row_index)
							s_index = col_alpha + '%d'%row_index
							obser_id = px.wb[sheet][s_index].value
							max_obser_list.append(obser_id)
							max_acc_v = str(round(acc_values[i],2))
							max_acc_values.append(max_acc_v)
						else:
							continue
					print("本次累计最大点:{},值:{}".format(\
						max_obser_list,max_acc_values))
			else:
					printl("warning, 没有最大累计值!")
					continue


			#累计变量报警值
			field_name = '累计变量报警值'
			field_values = self.get_values_by_field(sheet, row_list, field_name)
			print("DEBUG 累计变量报警值 数列:",field_values)
			s = ''
			if len(field_values) == 0: 
				s = ' '
			else:
				for value in field_values:
					if value == None:
						s += ' ' + '\n'
					else:
						s += str(value) + '\n'
			row.cells[6].text = s.strip('\n')


			#累计变化最大点
			if not self.alarm_feature:
				s = ''
				for obser in max_obser_list:
					s += obser + '\n' 
				row.cells[4].text = s.strip('\n')
			else:
				self.set_cell_text_by_field_values(row.cells[4], max_obser_list, max_acc_values, field_values)


			#累计变化率
			s = ''
			for max_acc_v in max_acc_values:
					s += max_acc_v + '\n'
			row.cells[5].text = s.strip('\n')	

			#累计变量控制值
			field_name = '累计变量控制值'
			field_values = self.get_values_by_field(sheet, row_list, field_name)
			s = ''
			if len(field_values) == 0: 
				s = ' '
			else:
				for value in field_values:
					if value == None:
						s += ' ' + '\n'
					else:
						s += str(value) + '\n'
			row.cells[7].text = s.strip('\n')


			#增加进度
			printl("%f@"%(sub_v_percent))
		#end for sheet in related_sheets


		#爆破振动 行
		row = t.add_row()
		row.cells[0].text = '爆破振动'
		second_cell = row.cells[1]
		second_cell.merge(row.cells[7])
		#巡检 行
		row = t.add_row()
		row.cells[0].text = '巡检'
		second_cell = row.cells[1]
		second_cell.merge(row.cells[7])
		row.cells[1].text = ''
		#数据分析 行
		row = t.add_row()
		row.cells[0].text = '数据分析'
		row.cells[1].merge(row.cells[7])
		row.cells[1].text = ''

		#设置表格样式
		t.alignment = WD_TABLE_ALIGNMENT.CENTER
		for row in t.rows:
			#设置高度:
			tr = row._tr
			trPr = tr.get_or_add_trPr()
			trHeight = OxmlElement('w:trHeight')
			trHeight.set(qn('w:val'), "420")
			trHeight.set(qn('w:hRule'), "atLeast")
			trPr.append(trHeight)

		#设置宽度，不起作用？
		t.autofit = False
		t.rows[0].width = Cm(2)
		for cell in t.rows[0].cells:
			cell.width = Cm(2)

		#设置字体，宋体
		ln = len(t.rows)
		last_three = len(t.rows) - 3
		for i in range(ln):
			for cell in t.rows[i].cells:
				for p in cell.paragraphs:
					p.style = d.styles["my_song_style"]
					if i >= last_three:
						p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


		#如果>0表示写了有效行的信息，否则返回False删除这个表，和表头
		if count_num > 0:
			return True,t
		else:
			return False,t
	##########one_overview_table()###############################


	def make_overview_pages(self):
		'''
		监测数据分析表
		'''
		result = False
		d = self.docx
		areas = self.my_xlsx.areas

		###new page###########
		d.add_page_break()
		p = d.add_paragraph()
		r = p.add_run("监测分析报告")
		r.font.size = Pt(15)
		r.bold = True
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		p.paragraph_format.space_before = 0
		p.paragraph_format.space_after = 0
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.paragraph_line_spacing = None

		p = d.add_paragraph()
		r = p.add_run("一、施工概况")
		r.bold = True
		r.font.size = Pt(14)

		for i in range(10):
			p = d.add_paragraph()
			r = p.add_run()
			r.font.size = Pt(12)

		###new page###########
		d.add_page_break()
		p = d.add_paragraph()
		r = p.add_run("二、数据分析")
		r.bold = True
		r.font.size = Pt(14)

		#表标题
		table_cap = "监测数据分析表"
		total_num = len(areas)
		count_num = 1
		v_percent = 12/total_num
		is_written = False
		for area_name in areas:
			#if '天目山路站' in area_name or '安薛区间' in area_name:
			if True:
				per_s = '[{}/{}]'.format(count_num, total_num)
				printl("[{}]数据分析表:'{}'".format(per_s, area_name))
				ss = '表' + '%d'%count_num + area_name + table_cap
				p = d.add_paragraph()
				p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
				p.paragraph_format.space_before = Pt(12)
				p.paragraph_format.space_after = Pt(2)
				r = p.add_run(ss)
				r.font.size = Pt(12)

				is_written,t = self.one_overview_table(area_name,v_percent,per_s)
				if not is_written:
					#删除表和表头
					delete_item(t)
					delete_item(p)
				else:
					count_num += 1

			#每写完两个表的时候，换页
			#最后一次不换页
			if count_num != 1 and (count_num-1) %2 == 0  and is_written:
				d.add_page_break()

		#最后一次如果不满偶数，也要换页
		if (count_num -1) %2 != 0:
			d.add_page_break()

		###new page###########
		#d.add_page_break()
		p = d.add_paragraph()
		r = p.add_run("三、结论")
		r.bold = True
		r.font.size = Pt(14)

		for i in range(5):
			p = d.add_paragraph()
			r = p.add_run()
			r.font.size = Pt(12)

		#表标题
		###new page###########
		d.add_page_break()
		p = d.add_paragraph()
		r = p.add_run("四、建议")
		r.bold = True
		r.font.size = Pt(14)

		for i in range(10):
			p = d.add_paragraph()
			r = p.add_run()
			r.font.size = Pt(12)

		ss = "监测单位：                   （盖章）              "
		p = d.add_paragraph()
		r = p.add_run(ss)
		r.font.size = Pt(14)
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		ss = "负责人：                   年     月     日        "
		p = d.add_paragraph()
		r = p.add_run(ss)
		r.font.size = Pt(14)
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

		result = True
		return result
	##################make_overview_pages()##############################


	def one_security_table(self, area_name):
		'''
		一个区间的现场巡查报表
		'''
		d = self.docx
		proj = self.proj
		ds = self.str_date

		t = d.add_table(rows=10, cols=6, style='security_table')
		t.cell(0,0).text = '线路名称'
		t.cell(0,1).text = proj.name
		t.cell(0,2).text = '监测标段'
		t.cell(0,3).text = ''
		t.cell(0,4).text = '工点名称'
		t.cell(0,5).text = area_name

		t.cell(1,0).text = '重点风险源'
		t.cell(1,1).merge(t.cell(1,3))
		t.cell(1,1).text = ''
		t.cell(1,4).text = '第三方监测单位'
		t.cell(1,5).text = proj.third_observer

		t.cell(2,0).text = '施工部位'
		t.cell(2,1).text = ''
		t.cell(2,2).text = '天气'
		t.cell(2,3).text = ''
		t.cell(2,4).text = '施工方监测单位'
		t.cell(2,5).text = proj.builder_observer

		t.cell(3,0).text = '巡视内容'
		t.cell(3,1).text = '存在的问题描述'
		t.cell(3,2).text = '原因分析'
		t.cell(3,3).text = '可能导致的后果'
		t.cell(3,4).text = '安全状态评价'
		t.cell(3,5).text = '处置措施建议'


		t.cell(4,0).text = '开挖面地质状况'
		t.cell(4,1).text = ''
		t.cell(4,2).text = '地质条件'
		t.cell(4,3).text = ''
		t.cell(4,4).text = ''
		t.cell(4,5).text = ''

		t.cell(5,0).text = '支护结构体系'
		t.cell(5,1).text = ''
		t.cell(5,2).text = ''
		t.cell(5,3).text = ''
		t.cell(5,4).text = ''
		t.cell(5,5).text = ''

		t.cell(6,0).text = '周边环境'
		t.cell(6,1).text = ''
		t.cell(6,2).text = ''
		t.cell(6,3).text = ''
		t.cell(6,4).text = ''
		t.cell(6,5).text = ''

		t.cell(7,0).text = '监测设施'
		t.cell(7,1).merge(t.cell(7,5))
		t.cell(7,1).text = ''

		t.cell(8,0).text = '现场巡视人'
		t.cell(8,1).merge(t.cell(8,2))
		t.cell(8,1).text = ' '*40+ ds
		t.cell(8,3).text = '项目技术负责人'
		t.cell(8,4).merge(t.cell(8,5))
		t.cell(8,4).text = ' '*40+ ds

		t.cell(9,0).merge(t.cell(9,5))
		s1 = "备注：1、本表由施工方和第三方监测单位采用；\n" 
		s2 = " "*12+"2、适用于XXX法施工；\n" 
		s3 = " "*12+"3、主要巡视内容包括：1）开挖面地质状况：土层性质及稳定性、降水效果和其它情况；"
		s4 = "支护结构体系：支护体系施作及时性、渗漏水情况、支护体系开裂、变形变化和其它情况；"
		s5 = "3）周边环境：建构筑物变形及开裂情况、地表变形及开裂情况、管线沿线地面开裂、渗水、塌陷情况、管线检查井开裂及积水变化和其它情况。"
		t.cell(9,0).text = s1+s2+s3+s4+s5

		#表格样式
		#加粗:
		for i in range(6):
			p = t.cell(3,i).paragraphs[0]
			for r in p.runs:
				r.font.bold = True
				r.font.size = Pt(12)

		for p in t.cell(9,0).paragraphs:
			p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
			for r in p.runs:
				r.font.bold = True
				r.font.size = Pt(10.5)

		#设置表格行高度
		for i in range(4,len(t.rows)):
			tr = t.rows[i]._tr
			trPr = tr.get_or_add_trPr()
			trHeight = OxmlElement('w:trHeight')
			v_height = "600"
			if i == len(t.rows) - 1:
				v_height = "1800"
			trHeight.set(qn('w:val'), v_height)
			trHeight.set(qn('w:hRule'), "atLeast")
			trPr.append(trHeight)

		#设置居中
		for i in range(len(t.rows)-1):
			row = t.rows[i]
			for cell in row.cells:
				for p in cell.paragraphs:
					p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
					for r in p.runs:
						r.font.size = Pt(12)


	##################one_security_table()###############################


	def make_security_pages(self):
		'''
		现场巡查报表
		'''
		result = False
		d = self.docx
		areas = self.my_xlsx.areas
		proj = self.proj

		table_cap = '现场巡查报表'
		i = 0
		total_num = len(areas)
		count_num = 0

		v_percent = 2/total_num
		for area_name in areas:
			count_num += 1
			#test debug only one area
			if True or '衡山路站' in area_name:
				printl("'[{}/{}]现场巡查报表:'{}'".format(count_num, total_num, area_name))
				i += 1
				ss = '表' + '%d'%i + ' 现场安全巡视表'
				p = d.add_paragraph()
				r = p.add_run(ss)
				r.font.bold = True
				r.font.size = Pt(12)
				p.paragraph_format.space_after = 0

				p = d.add_paragraph()
				p.add_run(area_name).underline = True
				p.add_run(table_cap)
				p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
				p.paragraph_format.space_after = 0
				for r in p.runs:
					r.font.bold = True
					r.font.size = Pt(16)

				p = d.add_paragraph()
				p.add_run('编号：')
				p.add_run(proj.code).underline = True
				p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
				p.paragraph_format.space_after = 0
				for r in p.runs:
					r.font.size = Pt(12)
				self.one_security_table(area_name)

				if count_num < total_num:
					d.add_page_break()

			printl("%f@"%v_percent)

		result = True
		return result
	#################make_security_pages()###############################


	def find_avail_rows_dates_values(self, sheet, area_name, needed_num, d_obser_range, obser_list):
		'''
		找到needed_num = 7天的有效值列
		返回三个列表:
		row_list = [row_index1, row_index2,...,row_indexy]
		date_list = [date7,date6,date5,date4,date3,date2,date1]
		value_list = [[date7_v1, date7_v2,...], [date6_v1, date6_v2,...],...] len(date_list) * len(row_list)
		'''
		px = self.my_xlsx
		row_list = []
		date_list = []
		value_list = []
		each_date_values = []

		#当天的有效行数index列表和值列表
		start_row, end_row = px.all_areas_row_range[sheet][area_name]
		#注意！list(range(3,5)) = [3,4], so need to add 1
		row_list = list(range(start_row, end_row+1))
		#获取当天日期的列坐标
		today_col_index, today_row_index = px.get_item_point(sheet, self.date)
		if today_col_index == None:
			return None, None, None

	
		today_values = self.get_col_values(sheet, area_name, today_col_index, \
			d_obser_range, obser_list)
		date_list.append(self.date)
		value_list.append(today_values)
		already_number = 1
		col_index = today_col_index
		#如果不够7天的数据，直到找到不为日期那一天为止
		while 1:
			col_index -= 1
			v_date = px.get_value(sheet, today_row_index, col_index)
			if not 'datetime' in str(type(v_date)):
				break
			date_list.append(v_date)
			lastday_values = self.get_col_values(sheet, area_name, col_index, \
				d_obser_range, obser_list)
			value_list.append(lastday_values)
			already_number += 1
			if already_number == needed_num:
				break

		return row_list, date_list, value_list
	##############find_avail_rows_dates_values()#############################################


	def draw_settlement_table(self, sheet, area_name, row_list, date_list, value_list,\
		 init_values, old_acc_values, cell_row=8):
		'''
		画沉降监测表格

		input:
		cell_row：该表格数据最大行数,
		'''
		d = self.docx
		px = self.my_xlsx

		col_num = 10
		if '支撑轴力' in sheet or '混撑' in sheet or '锚索轴力' in sheet:
			col_num = 12

		t = d.add_table(rows=13, cols=col_num, style='settlement_table')

		t_string = '沉降变化量(mm)'
		if '建筑物倾斜' in sheet:
			t_string = '位移变化量(‰)'
		elif '支撑轴力' in sheet or '混撑' in sheet or '锚索轴力' in sheet:
			t_string = '变化量(kN)'


		#日变量报警值
		field_name = '日变量报警值'
		diff_thresh_values = self.get_values_by_field(sheet, row_list, field_name)
		#累计变量报警值
		field_name = '累计变量报警值'
		acc_thresh_values = self.get_values_by_field(sheet, row_list, field_name)


		device_name = self.get_value_by_field(sheet, area_name, '仪器型号')
		device_code = self.get_value_by_field(sheet, area_name, '仪器出厂编号')
		check_date = self.get_value_by_field(sheet, area_name, '检定日期')
		s1 = "仪器型号：%s"%(device_name)
		s2 = " "*12 + "仪器出厂编号：%s"%(device_code)
		s3 = " "*12 + "检定日期：%s"%(check_date)

		if '支撑轴力' in sheet or '混撑' in sheet or '锚索轴力' in sheet:
			t.cell(0,0).merge(t.cell(0,11))
			t.cell(0,0).text = s1+s2+s3
			t.cell(1,0).merge(t.cell(2,0))
			t.cell(1,1).merge(t.cell(1,4))
			t.cell(1,0).text = '监测\n点号'
			t.cell(1,1).text = t_string
			t.cell(1,5).merge(t.cell(2,5))
			t.cell(1,5).text = '备注'
			t.cell(1,6).merge(t.cell(2,6))
			t.cell(1,6).text = '监测\n点号'
			t.cell(1,7).merge(t.cell(1,10))
			t.cell(1,7).text = t_string
			t.cell(1,11).merge(t.cell(2,11))
			t.cell(1,11).text = '备注'
	
			t.cell(2,1).text = '本次\n轴力'
			t.cell(2,2).text = '上次\n变量'
			t.cell(2,3).text = '本次\n变量'
			t.cell(2,4).text = '累计\n变量'
			t.cell(2,7).text = '本次\n轴力'
			t.cell(2,8).text = '上次\n变量'
			t.cell(2,9).text = '本次\n变量'
			t.cell(2,10).text = '累计\n变量'
		else:
			t.cell(0,0).merge(t.cell(0,9))
			t.cell(0,0).text = s1+s2+s3
			t.cell(1,0).merge(t.cell(2,0))
			t.cell(1,1).merge(t.cell(1,3))
			t.cell(1,0).text = '监测\n点号'
			t.cell(1,1).text = t_string
			t.cell(1,4).merge(t.cell(2,4))
			t.cell(1,4).text = '备注'
			t.cell(1,5).merge(t.cell(2,5))
			t.cell(1,5).text = '监测\n点号'
			t.cell(1,6).merge(t.cell(1,8))
			t.cell(1,6).text = t_string
			t.cell(1,9).merge(t.cell(2,9))
			t.cell(1,9).text = '备注'
	
			t.cell(2,1).text = '上次\n变量'
			t.cell(2,2).text = '本次\n变量'
			t.cell(2,3).text = '累计\n变量'
			t.cell(2,6).text = '上次\n变量'
			t.cell(2,7).text = '本次\n变量'
			t.cell(2,8).text = '累计\n变量'

		#填入数值
		#上次变量, 本次变量，累计量
		last_diffs = []
		this_diffs = []
		this_acc_diffs = []
		#本次轴力 
		this_values = list(map(lambda x:round(x,1),value_list[0]))
		this_values = array(this_values)
		#print("DEBUG 本次轴力:",this_values)

		ln_row = len(row_list)
		ln_date = len(date_list)
		#value_list = [[date7_v1, date7_v2,...], [date6_v1, date6_v2,...],...]
		#value_list should be ln_date*ln_row
		#init_values should be ln_row*1
		value_list = array(value_list, dtype=float)
		init_values = array(init_values, dtype=float)
		old_acc_values = array(old_acc_values, dtype=float)

		round_num = 2
		if '支撑轴力' in sheet or '混撑' in sheet or '锚索轴力' in sheet:
			round_num = 1

		if ln_date > 2:
			#今天和昨天差值 = 本次变量
			#today_diff_array =(value_list[0] - value_list[1])*1000
			today_diff_array = self.get_diff_values(sheet, value_list[0], value_list[1])
			today_diff = list(map(lambda x:round(x,round_num),today_diff_array))
			this_diffs = array(today_diff)

			#昨天和前天差值 = 上次变量
			#lastday_diff_array = (value_list[1] - value_list[2])*1000
			lastday_diff_array = self.get_diff_values(sheet, value_list[1], value_list[2])
			lastday_diff = list(map(lambda x:round(x,round_num),lastday_diff_array))
			last_diffs = array(lastday_diff)

			#今天和初值差值加旧累计 = 累计变量
			#today_acc_diff_array = (value_list[0] - init_values)*1000 + old_acc_values
			today_acc_diff_array = self.get_acc_values(sheet, value_list[0], init_values, old_acc_values)
			today_acc_diff = list(map(lambda x:round(x,round_num),today_acc_diff_array))
			this_acc_diffs = array(today_acc_diff)

		elif ln_date ==2:

			#今天和昨天差值 = 本次变量
			#today_diff_array =(value_list[0] - value_list[1])*1000
			today_diff_array = self.get_diff_values(sheet, value_list[0], value_list[1])
			today_diff = list(map(lambda x:round(x,round_num),today_diff_array))
			this_diffs = array(today_diff)

			#今天和初值差值加旧累计 = 累计变量
			#today_acc_diff_array = (value_list[0] - init_values)*1000 + old_acc_values
			today_acc_diff_array = self.get_acc_values(sheet, value_list[0], init_values, old_acc_values)
			today_acc_diff = list(map(lambda x:round(x,round_num),today_acc_diff_array))
			this_acc_diffs = array(today_acc_diff)

			#没有前天，上次变量设为'nan'
			last_diffs = array([None for x in range(ln_row)],dtype=float)

		elif ln_date == 1:
			#今天和初值差值加旧累计 = 累计变量
			#today_acc_diff_array = (value_list[0] - init_values)*1000 + old_acc_values
			today_acc_diff_array = self.get_acc_values(sheet, value_list[0], init_values, old_acc_values)
			today_acc_diff = list(map(lambda x:round(x,round_num),today_acc_diff_array))
			this_acc_diffs = array(today_acc_diff)

			this_diffs = array([None for x in range(ln_row)],dtype=float)
			last_diffs = array([None for x in range(ln_row)],dtype=float)

		else:
			printl("Error, date_list None")
			this_diffs = array([None for x in range(ln_row)],dtype=float)
			last_diffs = array([None for x in range(ln_row)],dtype=float)
			last_diffs = array([None for x in range(ln_row)],dtype=float)
		#print("DEBUG 沉降监测报表，本次累计: this_acc_diffs=",this_acc_diffs)

		obser_col,_ = px.get_item_point(sheet,'点号',from_last_search=False)
		if obser_col == None:
			obser_col,_ = px.get_item_point(sheet,'测点',from_last_search=False)
			if obser_col == None:
				print("Error, 没有点号列!")

		#表格变化值填写
		base_index = 3
		for i in range(cell_row):
			#如果观测点数小于cell行数，则当填写完观测点即退出
			if ln_row < cell_row and i == ln_row:
				break
			#监测点号, 注意表格格式，直接从第二列获取
			if obser_col != None:
				obser_name = px.get_value(sheet,row_list[i],obser_col)
			else:
				obser_name = 'Error'
			if '支撑轴力' in sheet or '混撑' in sheet or '锚索轴力' in sheet:
				#监测点号
				t.cell(base_index+i,0).text = obser_name
				r_name = t.cell(base_index+i,0).paragraphs[0].runs[0]
				#本次轴力:
				t.cell(base_index+i,1).text = str(this_values[i])
				#上次变量
				t.cell(base_index+i,2).text = str(last_diffs[i])
				#本次变量
				t.cell(base_index+i,3).text = str(this_diffs[i])
				if self.set_cell_text_by_field_value(t.cell(base_index+i,3),\
					this_diffs[i],diff_thresh_values[i]):
					r_name.bold = True
					#备注
					t.cell(base_index+i,5).text = '报警'
					t.cell(base_index+i,5).paragraphs[0].runs[0].bold=True

				#累计变量
				t.cell(base_index+i,4).text = str(this_acc_diffs[i])
				if self.set_cell_text_by_field_value(t.cell(base_index+i,4),\
					this_acc_diffs[i],acc_thresh_values[i]):
					r_name.bold = True
					#备注
					t.cell(base_index+i,5).text = '报警'
					t.cell(base_index+i,5).paragraphs[0].runs[0].bold=True

			else:
				#监测点号
				t.cell(base_index+i,0).text = obser_name
				r_name = t.cell(base_index+i,0).paragraphs[0].runs[0]
				#上次变量
				t.cell(base_index+i,1).text = str(last_diffs[i])
				#本次变量
				t.cell(base_index+i,2).text = str(this_diffs[i])
				if self.set_cell_text_by_field_value(t.cell(base_index+i,2),\
					this_diffs[i],diff_thresh_values[i]):
					r_name.bold = True
					#备注
					t.cell(base_index+i,4).text = '报警'
					t.cell(base_index+i,4).paragraphs[0].runs[0].bold=True

				#累计变量
				t.cell(base_index+i,3).text = str(this_acc_diffs[i])
				if self.set_cell_text_by_field_value(t.cell(base_index+i,3),\
					this_acc_diffs[i],acc_thresh_values[i]):
					r_name.bold = True
					#备注
					t.cell(base_index+i,4).text = '报警'
					t.cell(base_index+i,4).paragraphs[0].runs[0].bold=True


			#另外一侧的表格
			j = i+cell_row
			if ln_row > j:
				if obser_col != None:
					obser_name = px.get_value(sheet,row_list[j],obser_col)
				else:
					obser_name = 'Error'

				if '支撑轴力' in sheet or '混撑' in sheet or '锚索轴力' in sheet:
					t.cell(base_index+i,6).text = obser_name
					r_name = t.cell(base_index+i,6).paragraphs[0].runs[0]

					t.cell(base_index+i,7).text = str(this_values[j])
					t.cell(base_index+i,8).text = str(last_diffs[j])
					t.cell(base_index+i,9).text = str(this_diffs[j])
					if self.set_cell_text_by_field_value(t.cell(base_index+i,9),\
					this_diffs[j],diff_thresh_values[j]):
						r_name.bold = True
						#备注
						t.cell(base_index+i,11).text = '报警'
						t.cell(base_index+i,11).paragraphs[0].runs[0].bold=True



					t.cell(base_index+i,10).text = str(this_acc_diffs[j])
					if self.set_cell_text_by_field_value(t.cell(base_index+i,10),\
					this_acc_diffs[j],acc_thresh_values[j]):
						r_name.bold = True
						#备注
						t.cell(base_index+i,11).text = '报警'
						t.cell(base_index+i,11).paragraphs[0].runs[0].bold=True

				else:
					t.cell(base_index+i,5).text = obser_name
					r_name = t.cell(base_index+i,5).paragraphs[0].runs[0]

					t.cell(base_index+i,6).text = str(last_diffs[j])
					t.cell(base_index+i,7).text = str(this_diffs[j])
					if self.set_cell_text_by_field_value(t.cell(base_index+i,7),\
					this_diffs[j],diff_thresh_values[j]):
						r_name.bold = True
						#备注
						t.cell(base_index+i,9).text = '报警'
						t.cell(base_index+i,9).paragraphs[0].runs[0].bold=True


					t.cell(base_index+i,8).text = str(this_acc_diffs[j])
					if self.set_cell_text_by_field_value(t.cell(base_index+i,8),\
					this_acc_diffs[j],acc_thresh_values[j]):
						r_name.bold = True
						#备注
						t.cell(base_index+i,9).text = '报警'
						t.cell(base_index+i,9).paragraphs[0].runs[0].bold=True

		all_acc_diffs = []
		#all_acc_diffs = (array(value_list, dtype=float) - \
		#	array(init_values,dtype=float))*1000 + old_acc_values
		all_acc_diffs = self.get_acc_values(sheet, value_list, init_values, old_acc_values)


		t.cell(11,0).text = '累计变化量曲线图'
		if '支撑轴力' in sheet or '混撑' in sheet or '锚索轴力' in sheet:
			t.cell(11,1).merge(t.cell(11,11))
		else:
			t.cell(11,1).merge(t.cell(11,9))


		idx_list = []
		for row_idx in row_list:
			if obser_col != None:
				obser_name = px.get_value(sheet,row_idx,obser_col)
			else:
				obser_name = 'Error'
			idx_list.append(obser_name)

		#画图
		try:
			fig_path = self.my_plot.draw_settlement_fig(list(map(d_s,date_list)), \
				all_acc_diffs.transpose(), idx_list)
		except Exception as e:
			print("画图有问题: ",e)
		if fig_path == None or not os.path.exists(fig_path):
			printl("ERROR, fig_path not exists!",Flase)
			#fit_path = dummy.png
		else:
			#插入曲线图
			p = t.cell(11,1).paragraphs[0]
			run = p.add_run()
			run.add_picture(fig_path, width=Cm(12), height=Cm(5))
			p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		t.cell(12,0).text = '备注'
		if '支撑轴力' in sheet or '混撑' in sheet or '锚索轴力' in sheet:
			t.cell(12,1).merge(t.cell(12,11))
		else:
			t.cell(12,1).merge(t.cell(12,9))

		#填写 备注
		note = self.get_value_by_field(sheet, area_name, '备注')
		#t.cell(12,1).text = '1、“-”为下降、“+”为上升；2、监测点布设图见附图'
		t.cell(12,1).text = note

		#设置表格样式
		t.alignment = WD_TABLE_ALIGNMENT.CENTER
		for i in range(len(t.rows)):
			tr = t.rows[i]._tr
			trPr = tr.get_or_add_trPr()
			trHeight = OxmlElement('w:trHeight')
			v_height = "450"
			if i == 11:
				v_height = "3600"
			if i == 2:
				v_height = "600"
			trHeight.set(qn('w:val'), v_height)
			trHeight.set(qn('w:hRule'), "exact")
			trPr.append(trHeight)

			#中间观测点数据字体缩小
			if i >= 3 and i <= 10:
				for cell in t.rows[i].cells:
					for p in cell.paragraphs:
						for r in p.runs:
							r.font.size = Pt(8)

			#设置表格段落字体为自定义宋体样式
			for cell in t.rows[i].cells:
				for p in cell.paragraphs:
					p.style = d.styles["my_song_style"]

		#设置宽度
		#不起作用
		t.autofit = False
		t.rows[0].width = Cm(4.5)
		for cell in t.rows[0].cells:
			cell.width = Cm(4.5)
	##################draw_settlement_table()###################################


	def multi_settlement_table(self, area_name,v_percent,per_s):
		'''
		一个区间的多个沉降观测表

		步骤：
		找到初始值列，和邻近7天的观测值(包括当天)
		直到找到不为日期格式的列位置，有多少列观测有效值就添加多少列
		如果只有一列，即当天的，那么上次变化值为nan
		根据当天有效值的行数确定矩阵行数即为坐标的观测点行数范围，

		如果当天的值都为None，那么跳过该sheet.
		'''
		px = self.my_xlsx
		d = self.docx

		#找到这个area的所有观测项目
		related_sheets = []
		for sheet in px.sheets:
			if area_name in px.all_areas_row_range[sheet].keys():
					#related_sheets.append = [sheet1,sheet2,...]
					related_sheets.append(sheet)
		total_sheet_num = len(related_sheets)			
		print("{}个观测项目:{}".format(total_sheet_num, related_sheets))

		v_percent = v_percent/total_sheet_num
		count_num = 0
		#遍历这个站所有有关的测量数据,绘制表格	
		for sheet in related_sheets:
			count_num += 1
			if '测斜' in sheet:
				printl("%f@"%(v_percent))
				continue
			printl("{}沉降监测表:'{}'{}/{}({})".format(per_s, area_name,\
				count_num, total_sheet_num,sheet))
			#print("{}/{}'{}{}监测报表'".format(\
			#	count_num, total_sheet_num, area_name, sheet))

			#获取数据
			row_list = []   #观测点行坐标
			date_list = []  #7天日期 
			value_list = [] #7天有效数据2维列表
			initial_values = []
			old_acc_values = []

			#点号的行范围字典，适用于一个点号对应多行数值的sheet
			d_obser_range = {}
			obser_list = []

			obser_col,obser_row = px.get_item_point(sheet,'点号',from_last_search=False)
			if obser_col == None:
				obser_col,obser_row = px.get_item_point(sheet,'测点',from_last_search=False)
			if obser_col == None:
				printl("Error, {}无观测点列!".format(sheet))
				printl("%f@"%(v_percent))
				continue

			if '混撑' in sheet or '锚索轴力' in sheet :
				#以点号/测点为锚点，找每个点号/测点的行范围
				d_obser_range = px.get_one_sheet_areas_range(sheet, obser_col, obser_row+1)
				start, end = px.all_areas_row_range[sheet][area_name]
				for i in range(start, end+1):
					#获取第二列点号的值
					obser_name = px.get_value(sheet, i, obser_col)
					if obser_name != None:
						obser_list.append(obser_name)

			#找到该区间所有观测点的邻近7天的有效数据值
			row_list,date_list,value_list = \
			self.find_avail_rows_dates_values(sheet,area_name,7,d_obser_range,obser_list)
			#print("DEBUGdate_list=",date_list)
			if row_list == None:
				printl("Warning, 该区间'{}'在观测页'{}'没有{}当天值列!".format(\
					area_name,sheet,self.str_date))
				printl("%f@"%(v_percent))
				continue
			else:
				print("共{}/{}天有效观测数据".format(len(date_list),7))

			#初始值和旧累计列
			initial_name = '初值'
			if '混撑' in sheet:
				initial_name = '埋设前'

			init_col,_ = px.get_item_point(sheet, initial_name)
			if init_col == None:
				printl("Error, {}没有初值列!".format(sheet))
				printl("%f@"%(v_percent))
				continue
			else:
				initial_values = self.get_col_values(sheet, area_name, init_col, \
					d_obser_range,obser_list)
				if '混撑' in sheet:
					initial_values = -initial_values

			old_acc_col,_ = px.get_item_point(sheet, '旧累计')
			if old_acc_col != None:
				old_acc_values = self.get_col_values(sheet, area_name, old_acc_col, \
					d_obser_range,obser_list)
				#处理旧累计, 旧累计None的设为0
				ln_old_acc = len(old_acc_values)
				for i in range(ln_old_acc):
					if isnan(old_acc_values[i]):
						old_acc_values[i] = 0
			else:
				old_acc_values = [0 for x in range(len(initial_values))]


			###################################################
			#筛选保留当天值非nan的观测点, 其余nan的观测点不填写
			today_values = []
			today_col, _ = px.get_item_point(sheet, self.date)
			if today_col == None:
				printl("Error, 没有{}当天值列,略过该表!".format(self.date))
				continue

			today_values = self.get_col_values(sheet, area_name, today_col, \
				d_obser_range,obser_list)
			if isnan(today_values).sum() == len(today_values):
				printl("Warning, 当天没有有效值，略过该表!")
				continue

			c_row_list = deepcopy(row_list)
			c_value_list = deepcopy(value_list)
			c_initial_values = deepcopy(initial_values)
			c_old_acc_values = deepcopy(old_acc_values)

			#print("DEBUG today_values=",today_values)
			#print("DEBUG initial_values=",initial_values)
			#print("DEBUG old_acc_values=",old_acc_values)

			#筛选出非nan的index
			avail_index = []
			ln_today_values = len(today_values)
			for i in range(ln_today_values):
				if not isnan(today_values[i]):
					avail_index.append(i)

			#通过非nan的indiex，筛选保留非nan的数据
			row_list = [c_row_list[i] for _,i in enumerate(avail_index)]
			initial_values = [c_initial_values[i] for _,i in enumerate(avail_index)]
			old_acc_values = [c_old_acc_values[i] for _,i in enumerate(avail_index)]
			value_list = []
			for item in c_value_list:
				new_item = [item[i] for _,i in enumerate(avail_index)]
				value_list.append(new_item)

			#如果全部为nan，则略过这个表
			if len(row_list) == 0:
				print("所有点都没有有效值，略过该表")
				continue
			###################################################


			#计算每个表能填多少个观测点
			'''
			监测点号一边8个，共两边，按照总监测点是16的x倍数，
			则以总点数/x 来分
			x/16 > x//16:
			y = x//16+1
			or
			y = x//16
			'''
			ln = len(row_list)
			split_num = 0
			total_row = 16
			if ln/total_row > ln//total_row:
				split_num = ln//total_row + 1
			else:
				split_num = ln//total_row

			start = 0
			end = 0
			print("观测点数{}，共分{}组".format(ln,split_num))
			sub_v_percent = v_percent/split_num
			for i in range(1, split_num+1):
				#最后一个就是剩下的所有的
				if i == split_num:
					end = ln
				else:
					end = (ln//split_num)* i
				sub_row_list = row_list[start:end]
				#value_list = [[],[],[],...,[]] len(date) * len(rows)
				sub_value_list = [values[start:end] for values in value_list]
				sub_initial_values = initial_values[start:end]
				sub_old_acc_values = old_acc_values[start:end]
				start = end 

				#printl("'{}'{}监测报表{}/{}".format(\
				#	area_name, sheet,i,split_num))
				printl("{}沉降监测表:'{}'{}/{}({}{}/{})".format(per_s, area_name,\
					count_num, total_sheet_num,sheet,i,split_num))
				###new page###########
				if self.allow_page_break:
					d.add_page_break()
				else:
					self.allow_page_break = True

				#页头
				self.write_settlement_header(area_name)

				p = d.add_paragraph()	
				table_cap = area_name+sheet+'监测报表'+'%d/%d'%(i,split_num)
				r = p.add_run(table_cap)
				r.bold = True
				r.font.size = Pt(15)
				p.paragraph_format.space_before = Pt(6)
				p.paragraph_format.space_after = Pt(6)
				p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

				last_date = ''
				if len(date_list)==1:
					last_date = 'nan'
				else:
					last_date = date_to_str(date_list[1])
				p = d.add_paragraph()	
				p.add_run('上次监测时间：'+last_date)
				p.add_run(' '*34 + '本次监测时间：'+ self.str_date)
				for r in p.runs:
					r.font.size = Pt(11)
				p.paragraph_format.space_before = 0
				p.paragraph_format.space_after = 0
				p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
	
				#制表
				self.draw_settlement_table(sheet, area_name, sub_row_list, date_list,\
				 sub_value_list, sub_initial_values, sub_old_acc_values, total_row//2)

				#页尾
				self.write_settlement_foot(sheet, area_name)
				printl("%f@"%(sub_v_percent))
	#############multi_settlement_table()################################


	def write_settlement_header(self, area_name, show_area_name=True):
		'''
		沉降变化表/测斜表，项目信息
		'''
		d = self.docx
		p = d.add_paragraph()
		r = p.add_run(self.proj.name)
		r.font.size = Pt(16)
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		p.paragraph_format.space_before = 0
		p.paragraph_format.space_after = 0
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

		if show_area_name:
			p = d.add_paragraph()
			r = p.add_run("%s主体"%area_name)
			r.underline = True
			r.font.size = Pt(15)
			p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
			p.paragraph_format.space_before = 0
			p.paragraph_format.space_after = Pt(2)
			p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
			p.paragraph_format.line_spacing = Pt(26)

		p = d.add_paragraph()
		p.add_run("施工单位：")
		p.add_run(self.proj.builder).underline = True
		p.add_run(" "*12 + "编号：")
		p.add_run(self.proj.code).underline = True
		for r in p.runs:
			r.font.size = Pt(11)
		p.paragraph_format.space_before = 0
		p.paragraph_format.space_after = 0
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

		p = d.add_paragraph()
		p.add_run("监理单位：")
		p.add_run(self.proj.supervisor).underline = True
		for r in p.runs:
			r.font.size = Pt(11)
		p.paragraph_format.space_before = 0
		p.paragraph_format.space_after = 0
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

		p = d.add_paragraph()
		p.add_run("施工监测单位：")
		p.add_run(self.proj.builder_observer).underline = True
		for r in p.runs:
			r.font.size = Pt(11)
		p.paragraph_format.space_before = 0
		p.paragraph_format.space_after = 0
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

	################write_settlement_header()########################

	def insert_signature(self, r, name):
		'''
		插入签名图片
		'''
		#print("DEBUG file_list=",file_list)
		#print("添加签名",name)
		if len(self.sig_list) == 0:
			return

		for item in self.sig_list:
			sufx = os.path.basename(item)
			#print("DEBUG item=",sufx)
			if '.png' in sufx or '.PNG' in sufx:
				if name in sufx:
					try:
						r.add_picture(item, width=Cm(2), height=Cm(0.5))
						print('签名已添加:%s'%item)
					except Exception as e:
						pass
						print("Error: {}, item: {}".format(e,item))
					return True
		return False
	##########insert_signature()#####################################


	def write_settlement_foot(self, sheet, area_name):
		'''
		沉降变化表/测斜表页脚信息
		'''
		d = self.docx

		p = d.add_paragraph()
		person_name = self.get_value_by_field(sheet, area_name, '现场监测人')
		s = "现场监测人："
		r = p.add_run(s)
		if not self.insert_signature(r,person_name):
			r = p.add_run(person_name)


		person_name = self.get_value_by_field(sheet, area_name, '计算人')
		s = " "*6 + "计算人："
		r = p.add_run(s)
		if not self.insert_signature(r,person_name):
			r = p.add_run(person_name)

		person_name = self.get_value_by_field(sheet, area_name, '校核人')
		s = " "*6 + "校核人："
		r = p.add_run(s)
		if not self.insert_signature(r,person_name):
			r = p.add_run(person_name)

		for r in p.runs:
			r.font.size = Pt(11)
		p.paragraph_format.space_before = 0
		p.paragraph_format.space_after = 0

		p = d.add_paragraph()
		person_name = self.get_value_by_field(sheet, area_name, '监测项目负责人')
		s = "检测项目负责人："
		r = p.add_run(s)
		if not self.insert_signature(r,person_name):
			r = p.add_run(person_name)

		s = " "*6 + "第三方监测单位："
		p.add_run(s)
		p.add_run(self.proj.third_observer)
		for r in p.runs:
			r.font.size = Pt(11)
		p.paragraph_format.space_before = 0
		p.paragraph_format.space_after = 0
	##################write_settlementn_foot()###########################


	def make_settlement_pages(self):
		'''
		沉降变化监测表
		'''
		result = False
		d = self.docx
		areas = self.my_xlsx.areas
		proj = self.proj

		total_num = len(areas)
		count_num = 0
		per_s = ''
		#45 percent
		v_percent = 45/total_num
		for area_name in areas:
			count_num += 1
			#test debug only one area
			#if '衡山路站' in area_name:
			#if not '天目山路站' in area_name:
			#	continue
			if True:
				per_s = '[{}/{}]'.format(count_num,total_num)
				printl("{}沉降监测表:'{}'".format(per_s, area_name))		
				self.multi_settlement_table(area_name,v_percent, per_s)


		result = True
		return result

	################make_settlement_pages()##########################


	def one_inclinometer_table(self, sheet, area_name, sub_obser_list, d_obser_data, max_deep_values):
		'''
		一个测斜监测表, 含有两个或者一个观测点，按照深度的变化数据
		input:
		#d_obser_data = {'obser1':(deep_values,today_values, this_diffs, acc_diffs),'obser2':..}
		'''
		px = self.my_xlsx
		d = self.docx


		#画表
		t = d.add_table(rows=51, cols=9, style = 'Table Grid')
		t.cell(0,0).merge(t.cell(0,8))


		device_name = self.get_value_by_field(sheet, area_name, '仪器型号')
		device_code = self.get_value_by_field(sheet, area_name, '仪器出厂编号')
		check_date = self.get_value_by_field(sheet, area_name, '检定日期')
		s1 = "仪器型号：%s"%(device_name)
		s2 = " "*12 + "仪器出厂编号：%s"%(device_code)
		s3 = " "*12 + "检定日期：%s"%(check_date)
		t.cell(0,0).text = s1+s2+s3

		t.cell(1,0).text = '测点'
		t.cell(1,1).merge(t.cell(1,4))
		t.cell(1,5).merge(t.cell(1,8))
		t.cell(2,0).text = '孔深'
		t.cell(2,1).merge(t.cell(2,4))
		t.cell(2,5).merge(t.cell(2,8))

		t.cell(3,0).text = '深度(m)'
		for i in range(2):
			t.cell(3,1+4*i).text = '本次测值(mm)'
			t.cell(3,2+4*i).text = '本次变化(mm)'
			t.cell(3,3+4*i).text = '累计变化(mm)'
			t.cell(3,4+4*i).merge(t.cell(49,4+4*i))
		t.cell(50,0).merge(t.cell(50,8))
		#填写 说明
		note = self.get_value_by_field(sheet, area_name, '备注')
		t.cell(50,0).text = note 
		#s1 = '说明：1: 孔底起测; '
		#s2 = '2: "-"为向坑外倾斜，"+"为向坑内倾斜; '
		#s3 = '3: 日变化量报警值±2mm/d，累计变化量报警值±24mm.'
		#t.cell(50,0).text = s1+s2+s3 

		#填数据
		ln_deep = len(max_deep_values)
		for i in range(ln_deep):
			#深度值列
			t.cell(4+i,0).text = str(max_deep_values[i])

		for i in range(len(sub_obser_list)):
			deep_values = d_obser_data[sub_obser_list[i]][0]
			today_values = d_obser_data[sub_obser_list[i]][1]
			diff_values = d_obser_data[sub_obser_list[i]][2]
			acc_values = d_obser_data[sub_obser_list[i]][3]
			ln_value =len(deep_values)
			ln_diff = len(diff_values)
			ln_acc = len(acc_values)

			#观测点
			t.cell(1,1+i*4).text = sub_obser_list[i]
			#孔深
			t.cell(2,1+i*4).text = str(deep_values[-1])+'m'

			for j in range(ln_deep):
				if j < ln_value:
					#本次测值today_values
					t.cell(4+j,1+i*4).text = str(round(today_values[j],2))
				else:
					break
				if j < ln_diff:
					#本次变化this_diffs
					t.cell(4+j,2+i*4).text = str(round(diff_values[j],2))
				if j < ln_acc:
					#累计变化acc_diffs
					t.cell(4+j,3+i*4).text = str(round(acc_values[j],2))

			#画一个观测点的测斜变化量图
			'''
			print("DEBUG {},deep_values:{},today_values:{},diff_values:{},acc_values:{}".format(\
				sub_obser_list[i],deep_values,list(map(lambda x:round(x,2),list(today_values))),\
				list(map(lambda x:round(x,2),list(diff_values))),\
				list(map(lambda x:round(x,2),list(acc_values)))))
			'''
			fig_path = self.my_plot.draw_inclinometer_fig(deep_values,diff_values,acc_values)
			p = t.cell(3,4+4*i).paragraphs[0]
			run = p.add_run()
			run.add_picture(fig_path, width=Cm(2.5), \
				height=Cm(deep_values[-1])*0.68+2.11)
			p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		#end for i in range(len(sub_obser_list)):
		#表格样式，段落居中，字体数据为7.5磅
		ln = len(t.rows)
		for i in range(ln):
			for cell in t.rows[i].cells:
				for p in cell.paragraphs:
					p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
					if i >= 3:
						for r in p.runs:
							r.font.size = Pt(7.5)
					else:
						for r in p.runs:
							r.font.size = Pt(9)
							#r.bold = True
			#设置高度:
			tr = t.rows[i]._tr
			trPr = tr.get_or_add_trPr()
			trHeight = OxmlElement('w:trHeight')
			v_height = "185"
			if i == 3:
				v_height = '350'
			if i < 3:
				v_height = '240'
			trHeight.set(qn('w:val'), v_height)
			trHeight.set(qn('w:hRule'), "exact")
			trPr.append(trHeight)
	###########one_inclinometer_table()##############################


	def make_inclinometer_pages(self):
		'''
		测斜报表
		1. 找到index的深度范围
		2. 2个index一组做表，以深度最大的为第一列
		3. 根据变化值填表
		4. 画图
		'''
		px = self.my_xlsx
		d = self.docx

		#获取测斜sheet的所有观测点和其深度范围的字典:
		#d_obser_deeps={'observer1:(3,33),'observer2:(34:74),...}
		d_obser_deeps = {}
		inc_sheet = ''
		obser_col = None

		#找到孔深测斜表
		for sheet_name in px.sheets:
			if '孔深测斜' in sheet_name:
				inc_sheet = sheet_name
				#以第二列点号为锚点，找每个点号的深度范围
				obser_col,obser_row = px.get_item_point(sheet_name,'点号',from_last_search=False)
				if obser_col == None:
					obser_col,obser_row = px.get_item_point(sheet,'测点',from_last_search=False)
				#所有点号深度的行范围字典
				d_obser_deeps = px.get_one_sheet_areas_range(sheet_name,obser_col,obser_row+1)
				break

		if inc_sheet != '':
			print("所有观测点的深度行号:", d_obser_deeps)
		else:
			print("Error, 没有找到测斜观测sheet!")
			return True

		#获取区间和观测点的字典:
		#d_area_obser = {'area1':('observer1','observ2','observer3'..),..}
		d_area_obser = {}
		obser_list = []
		for area_name in px.all_areas_row_range[inc_sheet].keys():
			start, end = px.all_areas_row_range[inc_sheet][area_name]
			for i in range(start, end+1):
				#获取第二列点号的值
				obser_name = px.get_value(inc_sheet, i, obser_col)
				if obser_name != None:
					obser_list.append(obser_name)
				else:
					continue
			if len(obser_list)>0:
				d_area_obser[area_name] = obser_list[:]
				obser_list[:] = []
		total_area_num = len(d_area_obser.keys())
		print("孔深测斜表共有区间{}, 对应观测点:{}".format(total_area_num,d_area_obser))

		#找到当天日期,初值，旧累计所在的列坐标
		init_col, _ = px.get_item_point(inc_sheet, '初值', False)
		if init_col == None:
			printl("Error,{}初值列缺失!".format(inc_sheet))
			return False
		old_acc_col,_ = px.get_item_point(inc_sheet, '旧累计', False)
		if old_acc_col == None:
			printl("Error,{}旧累计列缺失!".format(inc_sheet))
			return False
		today_col,today_row = px.get_item_point(inc_sheet, self.date, True)
		if today_col == None:
			printl("Error,{},{}当天值列缺失!".format(inc_sheet, self.str_date))
			return False
		deep_col,_ = px.get_item_point(inc_sheet, '深度', False)
		if deep_col == None:
			printl("Error,{}深度值列缺失!".format(inc_sheet))
			return False

		#遍历每个区间制作多个测斜表
		count = 0
		count_num = 0
		v_percent = 35/len(d_area_obser.keys())
		for area_name in d_area_obser.keys():
			#if not '天目山路站' in area_name:
			#	continue
			count_num += 1
			count += 1
			obser_list = d_area_obser[area_name]
			printl("[{}/{}]测斜监测表:'{}'".format(count_num, total_area_num, \
				area_name))

			##############################
			#筛选，去掉当天值全是nan的点号
			c_obser_list = deepcopy(obser_list)

			for obser in c_obser_list:
				start_row, end_row = d_obser_deeps[obser]
				row_list = list(range(start_row, end_row+1))
				_, today_values = px.get_avail_rows_values(inc_sheet, row_list,\
						today_col, accept_none=False)
				if len(today_values) == 0:
					print("略过观测点:{}".format(obser))
					obser_list.remove(obser)
			##############################

			ln = len(obser_list)
			table_num = 0
			if ln/2 > ln//2:
				table_num = ln//2 + 1
			else:
				table_num = ln//2
			n = 1
			#两个观测点一组，进行制表

			sub_v_percent = v_percent/table_num
			for i in range(0,ln,2):

				#获取数据
				#d_obser_data = {'obser1':(deep_values,today_values, this_diffs, acc_diffs),'obser2':..}
				d_obser_data = {}
				obser_data = []
				deep_values = []
				today_values = []
				lastday_values = []
				init_values = []
				old_acc_values = []
				max_deep_values = []
				last_date = None
				this_diffs = []
				acc_diffs = []

				sub_obser_list = obser_list[i:i+2]
				printl("[{}/{}]测斜监测表:'{}'{}/{}:{}".format(count_num, total_area_num, \
				area_name,n,table_num,sub_obser_list))

				for obser in sub_obser_list:
					start_row, end_row = d_obser_deeps[obser]
					#注意！list(range(3,5)) = [3,4], so need to add 1
					#注意这个row_list是全深度的范围，不是当天有效值的范围
					#如果遇到当天某个深度没有填写，使用array的nan
					row_list = list(range(start_row, end_row+1))
		
					#找到这个观测点对应行数范围内深度的数据
					_, deep_values = px.get_avail_rows_values(inc_sheet, row_list,\
					deep_col, False)
					#找这两个观测点的最大深度
					if len(deep_values) > len(max_deep_values):
						max_deep_values = deep_values
					#初值
					_, init_values = px.get_avail_rows_values(inc_sheet, row_list,\
					init_col, False)
					if (len(init_values)!=len(row_list)) or (len(deep_values)!=\
						len(row_list)):
						#不可能发生，因为row_list的范围就是根据初值锚定的
						printl("Error! 测斜表初始值/深度值缺失！")
						print("DEBUG obser:{},row_list:{},init_values:{}".format(\
							obser,row_list,init_values))
						#defence
						pass

					#旧累计, true 可以接受None
					_, old_acc_values = px.get_avail_rows_values(inc_sheet, row_list,\
					old_acc_col, True)
					#处理旧累计，如果为None就设为0
					ln_old_acc = len(old_acc_values)
					for i in range(ln_old_acc):
						if old_acc_values[i] == None:
							old_acc_values[i] = 0

					#当天数据
					_, today_values = px.get_avail_rows_values(inc_sheet, row_list,\
					today_col, True)
					#寻找前一天数据
					last_date = px.get_value(inc_sheet, today_row, today_col-1)
					if 'datetime' in str(type(last_date)):
						_, lastday_values = px.get_avail_rows_values(inc_sheet, row_list,\
					today_col-1, True)
					#前面一列不是日期值，表示昨天值不存在
					else:
						lastday_values = None

					#使用array(list,dtype=float)来处理None值为nan
					this_diffs = array(today_values,dtype=float) - array(\
						lastday_values,dtype=float)
					acc_diffs = array(today_values,dtype=float) - array(\
						init_values,dtype=float) + array(old_acc_values,dtype=float)
					obser_data.append(array(deep_values,dtype=float))
					obser_data.append(array(today_values,dtype=float))
					obser_data.append(this_diffs)
					obser_data.append(acc_diffs)
					d_obser_data[obser] = obser_data
					obser_data = []
				#end获取数据 


				#页面头信息
				d.add_page_break()
				self.write_settlement_header(area_name, False)
				#表标题
				p = d.add_paragraph()	
				s = area_name+inc_sheet+"监测报表"+"%d/%d"%(n,table_num)
				r = p.add_run(s)
				r.bold = True
				r.font.size = Pt(14)
				p.paragraph_format.space_before = Pt(6)
				p.paragraph_format.space_after = Pt(0)
				p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

				last_date = ''
				if lastday_values == init_values:
					last_date = '初始值'
				elif lastday_values == None:
					last_date = '无上次监测'
				else:
					last_date = date_to_str(px.get_value(inc_sheet, today_row, today_col-1))

				p = d.add_paragraph()	
				p.add_run('上次监测时间：'+last_date)
				p.add_run(' '*34 + '本次监测时间：'+ self.str_date)
				for r in p.runs:
					r.font.size = Pt(11)
				p.paragraph_format.space_before = 0
				p.paragraph_format.space_after = 0
				p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

				#画表填值
				self.one_inclinometer_table(inc_sheet, area_name, sub_obser_list, d_obser_data,\
					max_deep_values)

				#页面尾信息
				#p = d.add_paragraph()
				self.write_settlement_foot(inc_sheet, area_name)
				#该区间第几个子表计数	
				n+=1
				printl("%f@"%(sub_v_percent))
			#end 两个观测点一组，进行制表
			#for i in range(0,ln,2):

		printl("[{}/{}],测斜监测表'{}'完成\n".format(count_num, total_area_num, \
				area_name))
		#end 遍历每个区间制作多个测斜表
		#for area_name in d_area_obser.keys():

		return True
	###############make_inclinometer_pages()#########################


	def make_blasting_pages(self):
		'''
		爆破振动监测报表
		'''
		d = self.docx
		px = self.my_xlsx

		p = d.add_paragraph()
		r = p.add_run(self.proj.name)
		r.font.size = Pt(15)
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		p = d.add_paragraph()
		r = p.add_run('爆破振动监测报表')
		r.font.size = Pt(15)
		r.bold = True
		p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

		p = d.add_paragraph()
		p.add_run("施工监测单位：")
		p.add_run(self.proj.builder_observer)
		p.add_run(" "*30 + "第三方监测单位：")
		p.add_run(self.proj.third_observer)
		for r in p.runs:
			r.font.size = Pt(12)
		p.paragraph_format.space_after = 0

		t = d.add_table(rows=8, cols=13, style='blasting_style')
		t.cell(0,0).merge(t.cell(0,12))
		s1 = "仪器型号："
		s2 = " "*30 + "仪器出厂编号： "
		s3 = " "*30 + "检定日期："
		t.cell(0,0).text = s1+s2+s3
		t.cell(1,0).merge(t.cell(3,0))
		t.cell(1,0).text = "测量时间"
		t.cell(1,1).merge(t.cell(3,1))
		t.cell(1,1).text = "爆破位置"
		t.cell(1,2).merge(t.cell(3,2))
		t.cell(1,2).text = "测量地点"
		t.cell(1,3).merge(t.cell(3,3))
		t.cell(1,3).text = "爆破中心至测点距离(m)"

		t.cell(1,4).merge(t.cell(1,5))
		t.cell(1,4).text = '爆破参数'
		t.cell(2,4).merge(t.cell(3,4))
		t.cell(2,4).text = "起爆药量(kg)"
		t.cell(2,5).merge(t.cell(3,5))
		t.cell(2,5).text = "段最大药量(kg)"

		t.cell(1,6).merge(t.cell(1,11))
		t.cell(1,6).text = "振动速度及主频频率"

		t.cell(2,6).merge(t.cell(2,7))
		t.cell(2,6).text = "最大向径分量"
		t.cell(2,8).merge(t.cell(2,9))
		t.cell(2,8).text = "最大切向分量"
		t.cell(2,10).merge(t.cell(2,11))
		t.cell(2,10).text = "最大垂直分量"


		t.cell(1,12).merge(t.cell(3,12))
		t.cell(1,12).text = "允许爆破振速度(cm/s)"

		t.cell(3,6).text = "v1(cm/s)"
		t.cell(3,7).text = "f1(Hz)"
		t.cell(3,8).text = "v2(cm/)"
		t.cell(3,9).text = "f(Hz)"
		t.cell(3,10).text = "v3(cm/s)"
		t.cell(3,11).text = "f3(Hz)"

		t.cell(7,0).text = "备注: "
		t.cell(7,1).merge(t.cell(7,12))

		p = d.add_paragraph()
		s = "现场监测人："
		p.add_run(s)
		s = " "*30 + "计算人："
		p.add_run(s)
		s = " "*30 + "校核人："
		p.add_run(s)
		s = " "*30 + "监测项目负责人："
		p.add_run(s)

		#表格样式，字体 宋体
		for row in t.rows:
			for cell in row.cells:
				for p in cell.paragraphs:
					p.style = d.styles["my_song_style"]
			#设置高度:
			tr = row._tr
			trPr = tr.get_or_add_trPr()
			trHeight = OxmlElement('w:trHeight')
			trHeight.set(qn('w:val'), "500")
			trHeight.set(qn('w:hRule'), "atLeast")
			trPr.append(trHeight)

		return True
	###############make_blasting_pages()########################


	def make_layout_pages(self):
		'''
		平面布点图
		把self.xlsx_path下的图片文件追加的docx中
		'''
		d = self.docx
		#获取文件夹下的所有文件地址:
		layout_path = os.path.join(self.xlsx_path,'平面布点图')
		file_list = os.listdir(layout_path)
		print("DEBUG 平面布点图文件:",file_list)
		for item in file_list:
			item_path = os.path.join(layout_path, item)
			if os.path.isfile(item_path) and ('.png' in item or '.PNG' in item or '.jpg' in item \
				or '.JPG' in item or 'jpeg' in item):
				try:
					d.add_picture(item_path, width=Cm(25), height=Cm(14))
					print("插入平面布点图: %s"%(item))
				except Exception as e:
					pass
					print("Error: {}, item: {}".format(e,item))

		return True
	#####################concatenate_new_docx()#######################

	def set_cell_text_by_field_value(self, cell, value, thresh):
		'''
		根据对比设置cell的text是否是bold
		True表示报警
		'''
		if self.alarm_feature:
			if 	not self.my_alarm.compare_threshold_safe(value,thresh):
					r = cell.paragraphs[0].runs[0]
					r.bold = True  
					return True

		return False
	##########set_cell_text_by_filed_value()#########################



	def set_cell_text_by_field_values(self, cell, cell_texts, cell_values, thresh_values):
		'''
		根据对比cell_values是否超出thresh_values
		对cell的数据进行bold
		'''
		ln_cell = len(cell_texts)
		ln_thresh = len(thresh_values)
		if ln_cell != ln_thresh:
			printl("Warning, 单元格值数目和报警值数目长度不匹配")

		p_cell = cell.paragraphs[0]
		r = None
		if ln_cell == 0: 
			printl("warning, 单元格没有值！")
			cell.text = ' '
		else:
			for i in range(ln_cell):
				if i != ln_cell-1:	
					if cell_texts[i] == None:
						r = p_cell.add_run(' \n')
					else:
						r = p_cell.add_run(cell_texts[i] + '\n')
				else:
					if cell_texts[i] == None:
						r = p_cell.add_run(' ')
					else:
						r = p_cell.add_run(cell_texts[i])
				if i < ln_thresh and thresh_values[i]:
					if 	not self.my_alarm.compare_threshold_safe(cell_values[i],\
						thresh_values[i]):
						print("达到报警点！cell值:'{}',报警值:'{}'".format(cell_values[i],thresh_values[i]))
						r.font.bold = True
		print("DEBUG, 填写单元格，并且对比报警值结束")
		return
	###############set_cell_text_by_field_values()################
########################class MyDocx()####################################

class MyAlarm(object):
	def __init__(self):
		self.alarm_on = False
	#############__init__()##############


	def compare_threshold_safe(self, value, threshold):
		'''
		比较报警值，如果大于大的或者小于小的就返回False
		安全就True
		'''
		is_safe = True

		def is_number(s):
			try:
				float(s)
				return True
			except ValueError:
				pass
			return False

		#print("DEBUG threshold=",threshold)
		if threshold == None or str(threshold).strip(' ') == '':
			return True
		if not is_number(value):
			print("Error，不是有效的数值:{}".format(value))
			return True
		else:
			value = float(value)

		threshold = str(threshold)
		min_thr = None
		max_thr = None
		#有两组值的情况:
		if '/' in threshold:
			v = threshold.strip(' ').split('/')
			min_thr, max_thr = v[0], v[1]
			if '+' in v[0]:
				max_thr = v[0]
			if '-' in v[1]:
				min_thr = v[1]

			if is_number(min_thr.strip('±')):
				min_thr = float(min_thr.strip('±'))
			else:
				min_thr = None
			if is_number(max_thr.strip('±')):
				max_thr = float(max_thr.strip('±'))
			else:
				max_thr = None
		else:
			if is_number(threshold.strip(' ').strip('±')):
				if '±' in threshold:
					min_thr = -float(threshold.strip(' ').strip('±'))
					max_thr = float(threshold.strip(' ').strip('±'))
	
				#只有一组值的情况,当做最大值来对待
				else:
					max_thr = float(threshold.strip(' '))
			else:
				return True
				#max_thr == None
				#min_thr == None

		print("DEBUG value: {}, min_thr: {}, max_thr: {}, threshold: '{}'".\
			format(value,min_thr,max_thr,threshold))

		if max_thr != None:
			if value >= max_thr:
				self.alarm_on = True
				return False

		if min_thr != None:
			if value <= min_thr:
				self.alarm_on = True
				return False

		return True
	#####compare_threshold_safe()#########	





############class MyAlarm()#####################################3


def thread_test():
	import threading
	import multiprocessing

	#main thread is not in mainloop Error
	t = threading.Thread(target=run_test)
	t.start()

	#p = multiprocessing.Process(target=run_test)
	#p.start()
	#run_test()


def run_test():

	print("Start Test")
	xlsx_path = r'C:\Users\tarzonz\Desktop\演示工程A\一二工区计算表2018.1.1_da.xlsx' 
	docx_path = r'C:\Users\tarzonz\Desktop\演示工程A\demo1.docx'
	date_v = '2018/1/1'
	date_v = datetime.strptime(date_v, '%Y/%m/%d')
	project_info = ["青岛市地铁1号线工程", "一、二工区", "DSFJC02-RB-594", \
	"M1-ZX-2016-222", "中国中铁隧道局、十局集团有限公司",\
	 "北京铁城建设监理有限责任公司", "中国铁路设计集团有限公司",\
	 '中铁隧道勘察设计研究院有限公司', xlsx_path, date_v]

	my_xlsx = read_xlsx.MyXlsx(xlsx_path)
	my_docx = MyDocx(docx_path, project_info, my_xlsx, True)

	res = my_docx.gen_docx()	
	if res:
		print("'{}' has been created".format(docx_path))
		print("Done")


if __name__ == '__main__':

	#测试
	#run_test()
	thread_test()
